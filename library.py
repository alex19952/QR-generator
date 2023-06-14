import sys
from threading import Thread as td
from tkinter import Tk as tkinter_Tk, messagebox as tkinter_messagebox
from tkinter.filedialog import asksaveasfile as tkinter_asksaveasfile
from io import BytesIO
from os import path as os_path, devnull as os_devnull, mkdir as os_mkdir
from datetime import date as datetime_date, datetime as datetime_datetime
from subprocess import Popen as subprocess_Popen
import GUI
import Templates

from docx.oxml import OxmlElement as docx_OxmlElement
from openpyxl import load_workbook as openpyxl_load_workbook
from qrcode import make as qrcode_make
from docx import Document
from docx.oxml import ns
from PyQt5 import Qt
from PyQt5.QtWidgets import (QApplication)


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os_path.abspath(".")
    path = os_path.join(base_path, relative_path)
    if '\\' in path:
        path = path.replace('\\', '/')
    
    return path


def startGUI(*args):
    f = open(os_devnull, 'w')
    sys.stdout = f

    #распакова аргументовpython3 --version
    # path_resourses = '\\'.join(args[0].split('\\')[:-1]) + '\\Resourses'
    path_resourses = resource_path("Resourses\\")
    path_resourses = path_resourses.replace('\\', '/')
    path_excel_file = args[1]
    list_rows_exl = list(map(int, args[2].split()))
    list_num_col_QR_descr = list(map(int, args[3].split()))

    #установка системной локали
    translator = Qt.QTranslator()
    locale = Qt.QLocale.system().name()
    translator.load('qtbase_%s' % locale, Qt.QLibraryInfo.location(Qt.QLibraryInfo.TranslationsPath))
    
    #создание и запуск GUI
    app = QApplication(sys.argv)
    app.installTranslator(translator)
    win = GUI.Window("Система инвентаризации", path_resourses)


    #запуск основной программы
    def startThreadMain():
        
        thread_main = td(name='main', target=startMain)
        thread_main.start()

    win.main_bytton.clicked.connect(startThreadMain)


    #основная программа
    def startMain():

        win.frame_menu.hide()
        win.frame_load.show()

        template_index = win.template_index

        list_template = [Templates.template_65, Templates.template_40, Templates.template_24, Templates.template_12]
        TemplClass = list_template[template_index]
        var = TemplClass()

        workbook = openpyxl_load_workbook(path_excel_file)
        excel = workbook.active
        document = Document()
        set_document_format(document, var)

        make_QR_doc_list(excel, list_rows_exl, document, list_num_col_QR_descr, var)

        # сохранение екселя в массив
        # с сортировкой по кабинетам
        if win.create_base_for_DCT:
            data = []
            data_row = []
            index_i = 2
            while True:
                if excel.cell(index_i, 1).value is None:
                    break
                
                for index_j in range(1, 11):
                    data_row.append(excel.cell(index_i, index_j).value)

                data.append(data_row)

                index_i += 1
                data_row = []

            # удаление пробелов
            for row in data:
                row[7] = str(row[7]).replace(' ', '').replace(',', '.')

            # сортировка
            def custom_key(list_): 
                return str(list_[7])
            data.sort(key=custom_key)

            # сохранение первой строки
            main_row = []
            for j in range(1, 11):
                main_row.append(excel.cell(1, j).value)

        ###

        win.frame_load.hide()
        win.frame_menu.show()

        root = tkinter_Tk()
        root.withdraw()
        if os_path.isfile(win.path_resourses + '/icon.ico'):
            root.iconbitmap(win.path_resourses + '/icon.ico')
        save = None
        while True:
            try:
                save = tkinter_asksaveasfile(title="Сохранить файл", defaultextension=".docx",
                                        filetypes=(("Документ Word", "*.docx"),), initialfile="QR коды")
            except Exception as exc:
                tkinter_messagebox.showerror(title="Ошибка", message="Нет доступа.\n\nВозможно, указанный файл открыт в другой программе.")
            else:
                break

        root.destroy()
        
        if save:
            save_path = save.name
            save_dir = '/'.join(save_path.split('/')[:-1])
            save_doc_path = save_path

            # сохранение массива с базой
            # в отдельные файлы по кабинетам
            date = str(datetime_date.today())
            hour = str(datetime_datetime.now().hour)
            minute = str(datetime_datetime.now().minute).rjust(2, '0')
            time_now = '-'.join([hour, minute])
            dt = '_'.join([date, time_now])
            if win.create_base_for_DCT:
                save_base_path = save_dir + '/base_' + dt + '/'
                if not os_path.isdir(save_base_path):
                    os_mkdir(save_base_path)
                num_room = ''
                room_base = None
                for row in data:
                    if str(row[7]) != num_room:
                        if room_base:
                            room_base.close()
                        num_room = str(row[7])
                        name_base = save_base_path + "base " + num_room + ".csv"
                        room_base = open(name_base, 'w', encoding='utf-8')
                        print(*main_row, file=room_base, sep=';')
                    print(*row, file=room_base, sep=';')
            ##

            workbook.save(save_doc_path + "base " + dt + ".xlsx")
            document.save(save_doc_path)
            win.path_doc_file = save_doc_path
            path = "explorer " + "\"" + save_dir + "\""
            path = path.replace('/', '\\')
            subprocess_Popen(path)
            win.flag_save = True


    def set_document_format(document, var):
        section = document.sections[0]
        section.page_height = var.page_height
        section.page_width = var.page_width
        section.top_margin = var.top_margin
        section.left_margin = var.left_margin
        section.right_margin = var.right_margin
        section.bottom_margin = var.bottom_margin
        document.styles['Table Grid'].font.size = var.font_size


    def set_table_format(table, var):
        table.autofit = False
        table.rows[0].height = var.height_row

        for index_qr_cell in range(var.quantity_columns * 2):
            cell = table.cell(0, index_qr_cell)
            if not(index_qr_cell % 2):
                set_cell_format(cell, var.width_cell_QR, 1, 1, var.indents_cells_QR)
            else:
                set_cell_format(cell, var.width_cell_descr, 1, 0, var.indents_cells_descr)

        borders = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
        for border in borders:
            node_borders = docx_OxmlElement('w:tblBorders')
            node_border = docx_OxmlElement("w:{}".format(border))
            node_border.set(ns.qn('w:val'), 'none')
            node_border.set(ns.qn('w:sz'), '0')
            node_borders.append(node_border)
            table._tbl.tblPr.append(node_borders)


    def set_cell_format(cell, width, v_align, align, indents_cell):
        cell.width = width
        cell.vertical_alignment = v_align
        cell.paragraphs[0].alignment = align
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = docx_OxmlElement('w:tcMar')
        for indent in indents_cell:
            node = docx_OxmlElement("w:{}".format(indent))
            node.set(ns.qn('w:w'), str(indents_cell.get(indent)))
            node.set(ns.qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)


    def make_QR_doc_list(excel, list_rows_excel, doc, list_col_QR_descr, var):
        win.gif.start()
        win.gif.setPaused(True)
        start_cell = win.num_start_cell
        start_row = (start_cell - 1) // var.quantity_columns
        start_column = (start_cell - 1) % var.quantity_columns
 
        #добавление пустых строк
        for i in range(start_row + 1):
            table = doc.add_table(1, var.quantity_columns * 2, style='Table Grid')
            set_table_format(table, var)
        count_columns = start_column * 2

        count = 0
        for num_row_excel in list_rows_excel: 
            #добавление таблицы(строки)
            if count_columns >= var.quantity_columns * 2:
                table = doc.add_table(1, var.quantity_columns * 2, style='Table Grid')
                set_table_format(table, var)
                count_columns = 0
        
            #добавление картинки
            run = table.rows[0].cells[count_columns].paragraphs[0].add_run()
            run.add_picture(make_img_QR(excel, num_row_excel), width = var.height_pic, height = var.height_pic)
            #добавление описания
            list_descr = []
            for i in list_col_QR_descr:
                if excel.cell(num_row_excel, i).value is None:
                    continue
                else:
                    data = '- ' + str(excel.cell(num_row_excel, i).value)
                    list_descr.append(data)

            for index, descr in enumerate(list_descr):
                if index < len(list_descr) - 1:
                    table.rows[0].cells[count_columns + 1].add_paragraph()
                paragraph = table.rows[0].cells[count_columns + 1].paragraphs[index]
                run = paragraph.add_run()
                run.add_text(descr)
            
            count_columns += 2
            count += 1
            win.gif.jumpToFrame(int(count / len(list_rows_excel) * 100))


    def make_img_QR(excel_file, num_row):
        data = str(excel_file.cell(num_row, 1).value)
        bytesIO = BytesIO()
        img = qrcode_make(data)
        img.save(bytesIO, format="png")
        return bytesIO

    sys.exit(app.exec_())
